import os
import json
import shutil
from datetime import datetime
from docx import Document
import anthropic

from config import ANTHROPIC_API_KEY, CLAUDE_MODEL, BASE_CV_PATH, GENERATED_CVS_DIR

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

# Keywords used to locate sections in the CV
SUMMARY_KEYWORDS    = ["summary", "profile", "professional profile", "objective"]
SKILLS_KEYWORDS     = ["core skill", "key skill", "competenc", "technical skill"]
EXPERIENCE_KEYWORDS = ["experience", "employment", "work history", "career"]


# ── DOCX section utilities ────────────────────────────────────────────────────

def _is_heading(para) -> bool:
    if para.style.name.startswith("Heading"):
        return True
    text = para.text.strip()
    if not text or len(text) > 80:
        return False
    return any(run.bold for run in para.runs if run.text.strip())


def _find_section(doc: Document, keywords: list) -> tuple | None:
    """
    Return (content_start, content_end) paragraph indices for the first
    section whose heading matches any keyword.
    """
    paras    = doc.paragraphs
    headings = [
        (i, p.text.strip().lower())
        for i, p in enumerate(paras)
        if _is_heading(p)
    ]
    for idx, (i, heading_text) in enumerate(headings):
        if any(kw in heading_text for kw in keywords):
            next_i = headings[idx + 1][0] if idx + 1 < len(headings) else len(paras)
            return (i + 1, next_i)
    return None


def _extract_section_text(doc: Document, keywords: list) -> str:
    result = _find_section(doc, keywords)
    if not result:
        return ""
    start, end = result
    return "\n".join(p.text for p in doc.paragraphs[start:end] if p.text.strip())


def _extract_full_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _set_para_text(para, text: str):
    """Replace a paragraph's text while preserving the formatting of its first run."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ""


def _update_section(doc: Document, keywords: list, new_text: str):
    """
    Overwrite the content paragraphs of a section with new_text.
    Handles multi-line new_text by distributing lines across existing paragraphs.
    """
    result = _find_section(doc, keywords)
    if not result or not new_text.strip():
        return

    start, end = result
    content_paras = [p for p in doc.paragraphs[start:end] if p.text.strip()]
    if not content_paras:
        return

    new_lines = [l.strip() for l in new_text.strip().splitlines() if l.strip()]

    for i, para in enumerate(content_paras):
        _set_para_text(para, new_lines[i] if i < len(new_lines) else "")


# ── JSON helper ───────────────────────────────────────────────────────────────

def _parse_json(text: str) -> dict:
    text = text.strip()
    if text.startswith("```"):
        text = text.split("```", 2)[1]
        if text.startswith("json"):
            text = text[4:]
        text = text.rsplit("```", 1)[0]
    return json.loads(text.strip())


def _safe_name(text: str) -> str:
    return "".join(c if c.isalnum() or c in " _-" else "_" for c in text).strip()


# ── Public API ────────────────────────────────────────────────────────────────

def generate_cover_letter(job: dict) -> str:
    """
    Generate a short, copy-paste-ready cover letter tailored to the job.
    Uses CV text for reference but does not modify any file.
    Returns the cover letter as a plain string.
    """
    doc     = Document(BASE_CV_PATH)
    cv_text = _extract_full_text(doc)

    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=512,
        messages=[{
            "role": "user",
            "content": f"""Write a short, professional cover letter for this QA engineer job application.

CANDIDATE CV:
{cv_text}

JOB DESCRIPTION:
{job['job_description']}

COMPANY: {job.get('company_name', '')}
ROLE: {job.get('job_title', '')}

Rules:
- 1-2 short paragraphs max
- First paragraph: why this specific role/company is interesting + 1-2 most relevant strengths backed by the CV
- Optional second paragraph: short closing with call to action (only if needed)
- Professional but natural tone — not robotic
- Do NOT include subject line, date, address headers, or signature
- Output the letter text only, nothing else""",
        }],
    )
    return response.content[0].text.strip()


def score_job(job: dict) -> dict:
    """
    Step 1 — cheap call: score + 2-sentence notes only.
    Sends CV as extracted text (no PDF upload needed).
    Returns: { match_score, match_notes }
    """
    doc     = Document(BASE_CV_PATH)
    cv_text = _extract_full_text(doc)

    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=256,
        messages=[{
            "role": "user",
            "content": f"""You are a recruiter screening a QA engineer's CV against a job.

CANDIDATE CV:
{cv_text}

JOB DESCRIPTION:
{job['job_description']}

Score how well this CV matches this job (0-100) and explain in 2 sentences.
Respond ONLY with valid JSON:
{{"match_score": <integer 0-100>, "match_notes": "<2 sentences>"}}""",
        }],
    )
    return _parse_json(response.content[0].text)


def rewrite_cv(job: dict) -> dict:
    """
    Step 2 — targeted rewrite: update ONLY Summary and Core Skills sections.

    Logic:
    - Read the candidate's Professional Experience for reference
    - Add to Summary any relevant keywords the candidate actually has experience with
    - Add missing relevant skills to Core Skills
    - Touch nothing else

    Returns: { cv_changes, cv_path }
    """
    doc = Document(BASE_CV_PATH)

    current_summary    = _extract_section_text(doc, SUMMARY_KEYWORDS)
    current_skills     = _extract_section_text(doc, SKILLS_KEYWORDS)
    experience_text    = _extract_section_text(doc, EXPERIENCE_KEYWORDS)

    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=1000,
        messages=[{
            "role": "user",
            "content": f"""You are a professional CV consultant helping a QA engineer tailor their CV for a specific role.

YOUR TASK:
- Update ONLY the Summary and Core Skills sections
- Do NOT change wording in Professional Experience or any other section
- Only add keywords/skills that are genuinely backed by the candidate's Professional Experience
- Preserve the original writing style and tone
- Keep the same structure/format as the original sections

JOB DESCRIPTION:
{job['job_description']}

CURRENT SUMMARY:
{current_summary}

CURRENT CORE SKILLS:
{current_skills}

PROFESSIONAL EXPERIENCE (read-only reference — do not quote or modify):
{experience_text}

Respond ONLY with valid JSON, no markdown fences:
{{
  "cv_changes": "<short bullet list of what was added/changed>",
  "new_summary": "<updated summary — same style, with relevant keywords woven in>",
  "new_skills": "<updated core skills — same format as original, with missing relevant skills added>"
}}""",
        }],
    )

    result = _parse_json(response.content[0].text)

    # Apply changes to a copy of the original DOCX
    doc = Document(BASE_CV_PATH)   # re-open clean copy
    _update_section(doc, SUMMARY_KEYWORDS, result["new_summary"])
    _update_section(doc, SKILLS_KEYWORDS,  result["new_skills"])

    company   = _safe_name(job.get("company_name", "Company"))
    title     = _safe_name(job.get("job_title",    "Role"))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    out_path  = os.path.join(GENERATED_CVS_DIR, f"CV_{company}_{title}_{timestamp}.docx")

    doc.save(out_path)
    result["cv_path"] = out_path
    return result
